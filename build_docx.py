"""
Build a professional Word document from DOCUMENTATION.md
Run: python build_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_spacing(para, before=0, after=6, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_spacing:
        from docx.shared import Pt as Pt2
        pf.line_spacing = Pt2(line_spacing)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B4C7E")
    pBdr.append(bottom)
    pPr.append(pBdr)

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"),   kwargs.get("val",   "single"))
        tag.set(qn("w:sz"),    kwargs.get("sz",    "4"))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get("color", "BDD7EE"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_inline_code(para, text):
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

def apply_inline_formatting(para, text):
    """Parse **bold**, *italic*, `code` inline markers and add formatted runs."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.font.name = "Calibri"
            r.font.size = Pt(11)
        if m.group(2):   # **bold**
            r = para.add_run(m.group(2))
            r.font.name  = "Calibri"
            r.font.size  = Pt(11)
            r.font.bold  = True
        elif m.group(3): # *italic*
            r = para.add_run(m.group(3))
            r.font.name   = "Calibri"
            r.font.size   = Pt(11)
            r.font.italic = True
        elif m.group(4): # `code`
            r = para.add_run(m.group(4))
            r.font.name  = "Courier New"
            r.font.size  = Pt(9.5)
            r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        pos = m.end()
    # trailing plain text
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.name = "Calibri"
        r.font.size = Pt(11)

# ── Document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Cover page ────────────────────────────────────────────────────────────────

doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run("RAN Optimization Dashboard")
r.font.name  = "Calibri"
r.font.size  = Pt(26)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
set_para_spacing(title_para, before=24, after=6)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle_para.add_run("AI-Powered Intent-Based Network Management\nfor 5G Heterogeneous Networks")
r.font.name   = "Calibri"
r.font.size   = Pt(14)
r.font.italic = True
r.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
set_para_spacing(subtitle_para, before=4, after=24)

add_horizontal_rule(doc)

# Abstract box
abs_para = doc.add_paragraph()
abs_para.paragraph_format.left_indent  = Cm(1.0)
abs_para.paragraph_format.right_indent = Cm(1.0)
set_para_spacing(abs_para, before=12, after=12)
r = abs_para.add_run("Abstract — ")
r.font.bold  = True
r.font.name  = "Calibri"
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
r = abs_para.add_run(
    "The RAN Optimization Dashboard is a full-stack, AI-powered web application that brings the principles "
    "of Intent-Based Networking (IBN) to life in a 5G Radio Access Network (RAN) environment. Network engineers "
    'express their operational goals in plain natural language -- such as "Prioritize emergency communications at '
    'the central hospital now" -- and a four-stage autonomous AI pipeline transforms that intent into a structured, '
    "standards-compliant network optimization plan. The system analyzes real network KPIs drawn from a 6G "
    "Heterogeneous Network dataset, generates 3GPP Release 18 configurations, visualizes affected cells on an "
    "interactive topology map, and quantifies the performance improvement through before/after KPI comparisons. "
    "This document covers the theoretical foundations, system architecture, implementation details, dataset "
    "requirements, API specification, and deployment procedures."
)
r.font.name  = "Calibri"
r.font.size  = Pt(10)
r.font.italic = True

add_horizontal_rule(doc)
doc.add_page_break()

# ── Section helper ────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=6)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(16)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2B4C7E")
    pBdr.append(bot)
    pPr.append(pBdr)

def h2(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=12, after=4)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)

def h3(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=3)
    r = p.add_run(text)
    r.font.name  = "Calibri"
    r.font.size  = Pt(11.5)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(0x40, 0x60, 0x8A)

def body(doc, text):
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=6)
    p.paragraph_format.first_line_indent = Pt(0)
    apply_inline_formatting(p, text)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    apply_inline_formatting(p, text)

def numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    apply_inline_formatting(p, text)

def code_block(doc, text):
    """Monospace shaded box."""
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name  = "Courier New"
    r.font.size  = Pt(8.5)
    r.font.color.rgb = RGBColor(0xE8, 0xF4, 0xFD)
    # shade background
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "0D1F35")
    pPr.append(shd)

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    set_para_spacing(p, before=4, after=4)
    r = p.add_run("Note: ")
    r.font.bold  = True
    r.font.name  = "Calibri"
    r.font.size  = Pt(10)
    r.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
    r2 = p.add_run(text)
    r2.font.name   = "Calibri"
    r2.font.size   = Pt(10)
    r2.font.italic = True

def figure_placeholder(doc, num, caption):
    p = doc.add_paragraph()
    set_para_spacing(p, before=8, after=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    # shaded box
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "EBF3FB")
    pPr.append(shd)
    r = p.add_run(f"[ FIGURE {num} ]  ")
    r.font.bold  = True
    r.font.size  = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)
    r2 = p.add_run(caption)
    r2.font.size   = Pt(9.5)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, "1F3A6E")
        set_cell_border(cell, color="1F3A6E")
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name  = "Calibri"
    # Data rows
    for ri, row_data in enumerate(rows):
        fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        row  = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            set_cell_border(cell, color="BDD7EE")
            cp = cell.paragraphs[0]
            r  = cp.add_run(str(val))
            r.font.size = Pt(9.5)
            r.font.name = "Calibri"
            if ci == 0:
                r.font.bold = True
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing after table

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — Introduction
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "1. Introduction")

body(doc,
    "The management of modern mobile networks is one of the most complex challenges in the "
    "telecommunications industry. A single 5G deployment may involve thousands of radio cells spread "
    "across a city, each continuously generating streams of performance data — throughput, latency, "
    "packet loss, signal quality — that operators must monitor and act upon in near real time [1]. "
    "When a major event fills a stadium, when a hospital requires ultra-reliable low-latency communication "
    "for remote surgery, or when a factory floor deploys hundreds of IoT sensors simultaneously, the network "
    "must adapt dynamically and precisely.")

body(doc,
    "Traditionally, this adaptation has required highly skilled radio frequency (RF) engineers to manually "
    "translate operational goals into low-level network configuration commands. This process is time-consuming, "
    "error-prone, and increasingly unsustainable as network density and complexity grow with each generation "
    "of wireless standards [4].")

body(doc,
    "The **RAN Optimization Dashboard** addresses this challenge by combining two transformative technologies: "
    "**Intent-Based Networking (IBN)** [3], which elevates the human-machine interface from command syntax to "
    "natural language goals, and **Agentic AI** [15], which deploys autonomous, specialized AI agents to reason "
    "over the intent and produce actionable network configurations grounded in real data.")

body(doc,
    "The result is a system where a network engineer types a single sentence and receives, within seconds, "
    "a complete optimization plan — including the specific cells affected, the 3GPP Release 18-compliant "
    "configuration parameters [11], and a projected KPI improvement — visualized on an interactive network map.")

figure_placeholder(doc, 1,
    "Full dashboard overview — Network Monitor page showing live topology, KPI cards, and time-series chart.\n"
    "Screenshot: http://localhost:4000")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — Background
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "2. Background and Theoretical Foundations")

h2(doc, "2.1  Telecommunications and the Evolution to 5G")
body(doc,
    "Telecommunications has undergone rapid generational evolution over the past four decades. Each generation "
    "introduced fundamentally new use cases beyond speed improvements [4]. **1G (1980s)** introduced analog voice; "
    "**2G (1990s)** digitized voice and added SMS and encryption (GSM); **3G (2000s)** enabled mobile data and "
    "video calling; **4G LTE (2010s)** delivered broadband mobile internet and the smartphone revolution. "
    "**5G (2020s)** is defined by three foundational service categories standardized by 3GPP [1]:")
for item in [
    "**eMBB (Enhanced Mobile Broadband):** Gigabit-class data rates for consumer applications.",
    "**URLLC (Ultra-Reliable Low-Latency Communication):** Sub-millisecond latency for mission-critical applications such as autonomous vehicles, remote surgery, and industrial automation.",
    "**mMTC (Massive Machine-Type Communication):** Connectivity for billions of IoT devices with minimal power consumption [14].",
]:
    bullet(doc, item)

figure_placeholder(doc, 2,
    "5G three-slice architecture diagram — eMBB, URLLC, mMTC bands with example use cases.\n"
    "Source: Block diagram (draw.io / Figma)")

h2(doc, "2.2  Radio Access Networks (RAN)")
body(doc,
    "The Radio Access Network (RAN) is the portion of a mobile network responsible for wirelessly connecting "
    "User Equipment (UE) to the core network — it is the 'last mile' of wireless connectivity and the primary "
    "site of network optimization work [7]. In 5G, base stations are called **gNodeBs (gNBs)**, each managing "
    "one or more radio cells. A cell handles radio scheduling, beamforming, handover management, and "
    "quality-of-service enforcement for all devices within its coverage zone.")
body(doc,
    "A key challenge in RAN management is the configuration space complexity. A 5G gNB running **NR "
    "(New Radio)** numerology can support multiple component carriers, MIMO layers, and network slices "
    "simultaneously [11]. Optimizing these dimensions for a specific operational goal requires deep domain "
    "expertise — precisely the expertise this system encodes into its AI agents.")

figure_placeholder(doc, 3,
    "5G RAN architecture diagram — UE → gNB → 5G Core Network signal path.\n"
    "Source: 3GPP / O-RAN documentation")

h2(doc, "2.3  Heterogeneous Networks (HetNet)")
body(doc,
    "A **Heterogeneous Network (HetNet)** combines multiple tiers of base stations with different "
    "transmission powers, coverage areas, and roles within the same geographic area [5]. Rather than "
    "relying solely on large macro cells, a HetNet deploys smaller cells as a dense underlay to handle "
    "high-traffic hotspots and improve spatial spectrum reuse. The four cell types in this system reflect "
    "the standard HetNet hierarchy:")

add_table(doc,
    ["Cell Type", "Coverage", "Typical Use Case"],
    [
        ["Macro",  "1–35 km",     "Wide-area coverage, rural and suburban"],
        ["Micro",  "200 m–2 km",  "Urban coverage, moderate capacity"],
        ["Pico",   "10–200 m",    "Indoor/outdoor hotspots, shopping malls"],
        ["Femto",  "< 10 m",      "Enterprise buildings, home offices"],
    ],
    col_widths=[3.5, 3.5, 8.5],
)

figure_placeholder(doc, 4,
    "HetNet topology — multi-tier cell visualization with hover tooltip showing cell type, zone, and KPIs.\n"
    "Screenshot: http://localhost:4000 (hover any cell node)")

h2(doc, "2.4  Key Performance Indicators in RAN")
body(doc,
    "Network engineers assess cell health through a standardized set of Key Performance Indicators (KPIs). "
    "The following metrics are monitored in this system:")
for item in [
    "**Throughput (Mbps):** Volume of user data transmitted per unit time. The most direct measure of user experience. Sustained values below 60 Mbps indicate significant degradation.",
    "**Latency (ms):** Round-trip packet delay. 5G URLLC targets sub-1 ms over-the-air latency. Values above 50 ms are a warning; above 80 ms is critical [1].",
    "**Resource Utilization (%):** Fraction of radio Resource Blocks (RBs) actively allocated. Values above 85% signal congestion risk.",
    "**Packet Loss Ratio:** Fraction of transmitted packets not delivered. Above 1% degrades TCP and video performance; above 5% indicates a serious fault.",
    "**Signal-to-Noise Ratio (SNR, dB):** Logarithmic ratio of signal to background noise. Higher SNR enables higher-order modulation (e.g., 256-QAM) and greater throughput.",
    "**QoS Satisfaction:** Composite score of how well a cell meets its 5QI (5G QoS Identifier) commitments across all active service classes [1].",
]:
    bullet(doc, item)

figure_placeholder(doc, 5,
    "KPI Summary Cards — live network-wide averages for throughput, latency, load, and cell health counts.\n"
    "Screenshot: http://localhost:4000")
figure_placeholder(doc, 6,
    "Live KPI Time-Series Chart — rolling 20-tick throughput, latency, and load history.\n"
    "Screenshot: http://localhost:4000 (after ~1 minute)")

h2(doc, "2.5  Intent-Based Networking (IBN)")
body(doc,
    "Intent-Based Networking (IBN) is a network management paradigm that shifts the operator interface "
    "from imperative commands to declarative goals [3]. The IBN concept was formalized by the IETF [3] and "
    "further developed by the TMForum [12] and ETSI [20] for autonomous network management. It represents "
    "the foundational philosophy behind 6G's vision of self-driving networks [12]. The translation from "
    "intent to action involves four steps that this system implements directly:")
for item in [
    "**Intent Recognition** — parsing natural language to identify the operational goal, affected zone, urgency, and network slice type.",
    "**Intent Validation** — checking that the requested goal is feasible given current network state.",
    "**Policy Generation** — translating the intent into specific network configuration parameters [2].",
    "**Execution and Verification** — applying the configuration and monitoring KPIs to confirm fulfillment (represented by the before/after KPI comparison).",
]:
    numbered(doc, item)

figure_placeholder(doc, 7,
    "IBN workflow diagram — Natural Language Intent → Intent Recognition → Policy Generation → Network Action.\n"
    "Source: Block diagram")
figure_placeholder(doc, 8,
    "Intent submission form — text area with sample intent typed in, and history chips below.\n"
    "Screenshot: http://localhost:4000/intent")

h2(doc, "2.6  Agentic AI and Multi-Agent Systems")
body(doc,
    "**Agentic AI** refers to AI systems that autonomously pursue goals by planning, making decisions, "
    "using tools, and executing multi-step tasks [10]. Rather than generating a single static response, "
    "an agentic system reasons over a goal, breaks it into sub-tasks, calls external tools, evaluates "
    "intermediate results, and iterates until the goal is achieved.")
body(doc,
    "A **Multi-Agent System (MAS)** deploys multiple specialized agents that collaborate to solve problems "
    "too complex for any single agent [10]. In this system, **CrewAI** [15] orchestrates a sequential crew "
    "where each agent receives all previous agents' outputs as context — building a progressively richer "
    "understanding before producing its output. Key design properties include:")
for item in [
    "**Specialization:** Each agent has a role, backstory, and task prompt focused on a specific domain.",
    "**Context chaining:** Later agents receive full structured outputs of earlier agents.",
    "**Tool use:** Agents invoke the `NetworkDataReader` CSV tool to ground reasoning in real data.",
    "**Separation of concerns:** Failures in one agent are handled without corrupting others.",
]:
    bullet(doc, item)

figure_placeholder(doc, 9,
    "Agent Pipeline Timeline — four numbered steps with checkmarks and per-agent summaries.\n"
    "Screenshot: http://localhost:4000/intent (after submitting any intent)")

h2(doc, "2.7  Large Language Models in Network Management")
body(doc,
    "**Large Language Models (LLMs)** are transformer-based neural networks [17] trained on massive "
    "text corpora, enabling them to understand and generate human language with remarkable reasoning "
    "capability. **Meta Llama 3.3 70B** [6] powers this system, offering natural language understanding, "
    "structured JSON output generation, 3GPP standards knowledge, and adaptive behavior across the full "
    "diversity of real-world intents — without requiring explicit rules for each scenario. Inference is "
    "provided by **Groq's** hardware-accelerated API [16], with the **LiteLLM** adapter normalizing "
    "provider calls inside CrewAI [15].")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — Architecture
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "3. System Architecture")

body(doc,
    "The system follows a clean client-server architecture with AI processing encapsulated in the backend. "
    "The Next.js [18] frontend communicates with the Flask backend through a single REST API endpoint.")

code_block(doc,
"Browser (Next.js)\n"
"  /              → Network Monitor  (live topology + KPIs)\n"
"  /intent        → Intent AI        (submit intent → results)\n"
"  /dataset       → Dataset Manager  (upload / inspect CSV)\n"
"         │ HTTP POST /api/intent\n"
"         ▼\n"
"Flask Backend (Python)\n"
"  server.py → ran_crew.py → CrewAI Sequential Crew\n"
"    Agent 1: Intent Parser   (NL text → intent JSON)\n"
"    Agent 2: RAN Planner     (intent → 3GPP config JSON)\n"
"    Agent 3: Network Monitor (CSV read by zone → KPI JSON)\n"
"    Agent 4: RAN Optimizer   (before/after KPIs → action)\n"
"  All agents: Groq Llama 3.3 70B via LiteLLM\n"
"         ▼\n"
"6G HetNet Dataset (CSV)\n"
"  5,000 rows · 50 cells · Location_Tag per cell\n"
"  Zones: hospital | stadium | factory | downtown | residential")

body(doc, "**Intent submission data flow:**")
for step in [
    "Engineer types a natural language intent in the browser and submits the form.",
    "Frontend sends POST /api/intent to the Flask backend.",
    "Backend instantiates the CrewAI sequential crew and starts Agent 1.",
    "Each agent executes in turn, passing its structured JSON output to the next as context.",
    "Agents 3 and 4 call the NetworkDataReader tool to retrieve real cell KPI data filtered by location zone.",
    "Backend assembles four agent outputs into a single result JSON and returns it.",
    "Frontend renders the result card: banner, agent timeline, KPI comparison, config table, highlighted topology.",
]:
    numbered(doc, step)

figure_placeholder(doc, 10,
    "End-to-end data flow diagram — Engineer intent → POST request → Flask → CrewAI agents → CSV → JSON response → result card.\n"
    "Source: Swimlane or flow diagram")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Tech Stack
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "4. Technology Stack")

h2(doc, "4.1  Backend")
add_table(doc,
    ["Technology", "Version", "Purpose"],
    [
        ["Python",         "3.11+",   "Core language with mature AI/data science ecosystem."],
        ["Flask",          "3.x",     "Lightweight WSGI REST API framework."],
        ["Flask-CORS",     "4.x",     "Cross-origin resource sharing for browser calls."],
        ["CrewAI",         "0.70+",   "Multi-agent orchestration framework [15]."],
        ["Groq API",       "—",       "Hardware-accelerated Llama 3.3 70B inference [16]."],
        ["LiteLLM",        "—",       "Universal LLM adapter used internally by CrewAI."],
        ["pandas",         "2.x",     "CSV dataset loading, filtering, and sampling [19]."],
        ["pydantic",       "2.x",     "Tool input schema validation."],
        ["python-dotenv",  "1.x",     "Environment variable management."],
        ["gunicorn",       "—",       "Production WSGI server for Railway deployment."],
    ],
    col_widths=[3.5, 2.5, 9.5],
)

h2(doc, "4.2  Frontend")
add_table(doc,
    ["Technology", "Version", "Purpose"],
    [
        ["Next.js",       "14.x",    "React meta-framework with App Router and file-based routing [18]."],
        ["TypeScript",    "5.x",     "Type-safe JavaScript, catches CSV-to-component schema mismatches at compile time."],
        ["Tailwind CSS",  "3.x",     "Utility-first CSS for consistent styling."],
        ["Recharts",      "2.x",     "D3-based React charting library for live KPI charts."],
        ["PapaParse",     "5.x",     "In-browser CSV parser for client-side dataset loading."],
        ["Lucide React",  "0.400+",  "Consistent icon library."],
    ],
    col_widths=[3.5, 2.5, 9.5],
)

h2(doc, "4.3  Deployment")
add_table(doc,
    ["Service", "Purpose"],
    [
        ["Railway", "Python/Flask backend hosting with Procfile auto-detection."],
        ["Vercel",  "Next.js frontend with global CDN and automatic Git deployments."],
    ],
    col_widths=[4.0, 11.5],
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Project Structure
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "5. Project Structure")

code_block(doc,
"ran-dashboard-backend/\n"
"├── server.py            # Flask app — routes, CORS, entry point\n"
"├── ran_crew.py          # CrewAI agents, tasks, pipeline logic\n"
"├── csv_tool.py          # NetworkDataReader tool (pandas)\n"
"├── fallbacks.py         # Rule-based fallback responses\n"
"├── requirements.txt     # Python dependencies\n"
"├── Procfile             # gunicorn start command\n"
"├── data/\n"
"│   └── 6G_HetNet_with_location.csv\n"
"├── docs/figures/        # Screenshots and diagrams\n"
"└── frontend/\n"
"    ├── app/\n"
"    │   ├── page.tsx           # Network Monitor (/)\n"
"    │   ├── intent/page.tsx    # Intent AI (/intent)\n"
"    │   └── dataset/page.tsx   # Dataset Manager (/dataset)\n"
"    ├── components/\n"
"    │   ├── network/NetworkTopology.tsx\n"
"    │   ├── network/KPICards.tsx\n"
"    │   ├── network/LiveChart.tsx\n"
"    │   ├── intent/IntentPanel.tsx\n"
"    │   ├── intent/IntentResultCard.tsx\n"
"    │   └── dataset/DatasetManager.tsx\n"
"    ├── hooks/useDataset.ts\n"
"    ├── hooks/useNetworkPoll.ts\n"
"    └── lib/types.ts  |  network.ts  |  api.ts")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Setup
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "6. Setup and Installation")

h2(doc, "6.1  Prerequisites")
for item in [
    "**Python 3.11+** — required for modern typing and match statements.",
    "**Node.js 18+** — required by Next.js 14.",
    "**Groq API key** — free at https://console.groq.com (provides Llama 3.3 70B access).",
]:
    bullet(doc, item)

h2(doc, "6.2  Backend Setup")
code_block(doc,
"cd ran-dashboard-backend\n"
"pip install -r requirements.txt\n"
"# Add GROQ_API_KEY=gsk_... to .env\n"
"python download_dataset.py\n"
"python server.py\n"
"# Runs at http://localhost:8000")

h2(doc, "6.3  Frontend Setup")
code_block(doc,
"cd frontend\n"
"npm install\n"
"# Optional: create frontend/.env.local\n"
"# NEXT_PUBLIC_BACKEND_URL=http://localhost:8000\n"
"npm run dev -- --port 4000\n"
"# Runs at http://localhost:4000")

note(doc, "If ports 3000/3001 are occupied, use --port 4000 or any available port.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — Dataset
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "7. Dataset Requirements")

body(doc,
    "The system enforces strict column validation in both the Python backend (before LLM processing) "
    "and the TypeScript frontend (before visualization). A dataset missing any required column is rejected "
    "with a descriptive error message — there is no silent fallback to synthetic data.")

h2(doc, "7.1  Required Columns")
add_table(doc,
    ["Column", "Type", "Description"],
    [
        ["Cell_ID",                   "integer", "Primary key linking dataset rows to topology nodes."],
        ["Cell_Type",                 "string",  "HetNet tier: Macro / Micro / Pico / Femto."],
        ["Achieved_Throughput_Mbps",  "float",   "Downlink throughput in Mbps. Primary user experience metric."],
        ["Network_Latency_ms",        "float",   "Round-trip packet delay in ms. Critical for URLLC."],
        ["Resource_Utilization",      "float",   "Fraction of RBs allocated (0–1). Load and congestion indicator."],
        ["Packet_Loss_Ratio",         "float",   "Fraction of packets lost (0–1). High values indicate overload."],
        ["Signal_to_Noise_Ratio_dB",  "float",   "Received signal quality in dB. Determines modulation order."],
        ["Bandwidth_MHz",             "float",   "Channel bandwidth. Determines maximum throughput capacity."],
        ["Location_X",                "float",   "X coordinate in coverage grid (0–1000). Topology positioning."],
        ["Location_Y",                "float",   "Y coordinate in coverage grid (0–1000). Topology positioning."],
        ["Location_Tag",              "string",  "Zone label: hospital | stadium | factory | downtown | residential. Required for intent routing."],
    ],
    col_widths=[4.5, 2.0, 9.0],
)

h2(doc, "7.2  Location Zone Mapping")
body(doc,
    "The default 6G HetNet dataset assigns Location_Tag values based on X/Y coordinates within "
    "a 1000×1000 normalized coverage grid, reflecting realistic urban deployment scenarios:")
add_table(doc,
    ["Zone", "X Range", "Y Range", "Network Use Case"],
    [
        ["hospital",    "0–350",    "0–350",    "URLLC — emergency services, remote patient monitoring."],
        ["factory",     "650–1000", "0–350",    "mMTC/URLLC — industrial IoT, robotic automation."],
        ["stadium",     "650–1000", "650–1000", "eMBB — mass gatherings, live events."],
        ["residential", "0–350",    "650–1000", "eMBB — home broadband, consumer IoT."],
        ["downtown",    "350–650",  "350–650",  "Mixed — dense urban, enterprise, pedestrian users."],
    ],
    col_widths=[3.0, 2.5, 2.5, 7.5],
)

figure_placeholder(doc, 11,
    "Location zone map — 1000×1000 grid divided into five colored zones with X/Y range labels.\n"
    "Source: Block diagram")
figure_placeholder(doc, 12,
    "Dataset Manager page — metadata view and 11-column requirements list.\n"
    "Screenshot: http://localhost:4000/dataset")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — AI Pipeline
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "8. AI Pipeline — The Four-Agent Crew")

body(doc,
    "The pipeline is defined in `ran_crew.py` and orchestrated by CrewAI [15] as a sequential crew — "
    "each agent executes in order and every output is passed as context to the next. All agents use "
    "Groq Llama 3.3 70B [6] via LiteLLM. If the pipeline fails, the system falls back to `fallbacks.py`, "
    "a deterministic rule-based engine. The frontend labels fallback responses with a warning badge.")

figure_placeholder(doc, 13,
    "Full Intent Result Card — complete AI pipeline output for a hospital emergency intent.\n"
    "Screenshot: http://localhost:4000/intent (after submitting 'Prioritize emergency communications at the central hospital')")

h2(doc, "8.1  Agent 1 — Intent Parser")
body(doc,
    "**Role:** Senior Telecom Intent Analyst. Transforms free-form engineer text into a machine-readable "
    "intent JSON structure. Classifies the intent into one of eight operational categories and maps it to "
    "the corresponding 3GPP network slice type [1], [8]:")
add_table(doc,
    ["Intent Type", "Description", "Slice"],
    [
        ["emergency",            "Medical, disaster response, or public safety",   "URLLC"],
        ["stadium_event",        "Large venue, sports event, concert",             "eMBB"],
        ["iot_deployment",       "Sensor networks, smart city infrastructure",     "mMTC"],
        ["industrial_iot",       "Factory automation, robotics",                   "mMTC / URLLC"],
        ["video_streaming",      "High-bandwidth media delivery",                  "eMBB"],
        ["smart_city",           "Traffic, utilities, urban management",           "mMTC"],
        ["voice_priority",       "VoIP, emergency calling",                        "URLLC"],
        ["general_optimization", "Baseline performance improvement",               "eMBB"],
    ],
    col_widths=[4.5, 7.0, 4.0],
)
figure_placeholder(doc, 14,
    "Intent Summary Banner — parsed intent type, slice type, zone, confidence %, and health score.\n"
    "Screenshot: close-up of the result card banner")

h2(doc, "8.2  Agent 2 — RAN Planner")
body(doc,
    "**Role:** Expert 5G Network Configuration Architect. Translates the intent JSON into a complete "
    "3GPP Release 18 [1], [11] network configuration covering:")
for item in [
    "**Network Slice parameters:** SST, slice differentiator, priority, preemption capability.",
    "**QoS parameters:** 5QI identifier, ARP, GBR, MBR, Packet Delay Budget [13].",
    "**RAN parameters:** NR numerology (subcarrier spacing μ), MIMO layers, DL/UL scheduler, carrier aggregation.",
]:
    bullet(doc, item)
body(doc,
    "For example, a hospital emergency intent produces URLLC with 5QI=1 (GBR, 100 ms PDB), "
    "high ARP priority, numerology μ=3 (120 kHz subcarrier spacing for minimal latency), "
    "and a round-robin scheduler.")
figure_placeholder(doc, 15,
    "Generated 3GPP Network Configuration Table — slice, QoS, and RAN parameters with values.\n"
    "Screenshot: close-up of the config section in the result card")

h2(doc, "8.3  Agent 3 — Network Monitor")
body(doc,
    "**Role:** 5G NOC Analyst. Calls the `NetworkDataReader` tool to fetch actual KPI measurements "
    "from a real cell in the relevant location zone, then performs a structured health assessment "
    "against standard thresholds. Location inference maps intent keywords to dataset zones:")
add_table(doc,
    ["Keywords in Intent", "Inferred Zone"],
    [
        ["hospital, emergency, clinic, ambulance, medical", "hospital"],
        ["stadium, concert, festival, match, arena, fans",  "stadium"],
        ["factory, manufacturing, industrial, robots",       "factory"],
        ["downtown, city center, CBD, urban core",           "downtown"],
        ["residential, suburb, home, neighbourhood",         "residential"],
    ],
    col_widths=[9.0, 6.5],
)
body(doc,
    "A composite health score (0–100) is computed: above 75 is healthy, 50–75 is warning, "
    "below 50 is critical.")

h2(doc, "8.4  Agent 4 — RAN Optimizer")
body(doc,
    "**Role:** Senior RAN Optimization Engineer. Synthesizes all prior outputs to select the most "
    "appropriate optimization action and quantify its impact by calling `NetworkDataReader` twice — "
    "before and after states [9]:")
add_table(doc,
    ["Network Condition", "Selected Action", "Description"],
    [
        ["Cell load > 80% + low throughput", "scale_bandwidth",  "Widen channel or add carrier component."],
        ["Cell load > 80%",                  "activate_cell",    "Bring adjacent cell online to offload traffic."],
        ["Latency or packet loss violation",  "modify_qos",       "Raise priority, tighten delay budget."],
        ["Slice type mismatch",               "adjust_priority",  "Reconfigure slice parameters."],
        ["Health score > 85, all healthy",    "energy_saving",    "Reduce power or sleep idle cells."],
    ],
    col_widths=[5.5, 3.5, 6.5],
)
body(doc,
    "**Before/After KPI Measurement:** The optimizer reads a current (degraded) cell state, then a "
    "healthy cell state as a post-optimization proxy. Improvement is calculated as "
    "|((after − before) / before)| × 100%, with direction interpreted per metric semantics.")
figure_placeholder(doc, 16,
    "KPI Before/After Comparison — Throughput, Latency, Resource Load, and Packet Loss cards with % improvement badges.\n"
    "Screenshot: close-up of the KPI comparison section")
figure_placeholder(doc, 17,
    "Affected Topology — zone-highlighted cells lit, others dimmed to 15% opacity, monitored cell in cyan.\n"
    "Screenshot: close-up of the topology section in the result card")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — API Reference
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "9. API Reference")

h2(doc, "9.1  GET /health")
body(doc, "Returns server status and model configuration.")
code_block(doc,
'{\n'
'  "status": "ok",\n'
'  "service": "RAN Optimizer Backend (CrewAI)",\n'
'  "groq_configured": true,\n'
'  "model": "llama-3.3-70b-versatile",\n'
'  "agents": ["Intent Parser","RAN Planner","Network Monitor","RAN Optimizer"]\n'
'}')

h2(doc, "9.2  POST /api/intent")
body(doc, "Runs the full four-agent pipeline. Typical processing time: 10–30 seconds.")
code_block(doc,
'// Request\n'
'{ "intent": "Prioritize emergency communications at the central hospital now" }\n\n'
'// Response — AI success\n'
'{ "success": true, "result": {\n'
'    "intent":       { "intent_type": "emergency", "slice_type": "URLLC", ... },\n'
'    "config":       { "network_slice": {...}, "qos_parameters": {...} },\n'
'    "monitor":      { "cell_id": 7, "location_tag": "hospital", ... },\n'
'    "optimization": { "action": "modify_qos", "before": {...}, "after": {...} }\n'
'}}\n\n'
'// Response — fallback\n'
'{ "success": true, "result": {...},\n'
'  "warning": "AI agents encountered an error. Showing rule-based fallback." }')

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — Frontend Pages
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "10. Frontend Pages")

h2(doc, "10.1  Network Monitor (/)")
body(doc,
    "Provides a continuously updated view of all 50 cells by advancing dataset rows at ~3-second "
    "intervals with small KPI jitter — simulating live operations center conditions. Four panels:")
for item in [
    "**KPI Summary Cards** — network-wide averages for throughput, latency, load, and health counts.",
    "**Network Topology** — SVG map of all cells. Color = health status. Hover shows tooltip with all KPIs and location zone.",
    "**Live Chart** — Recharts time-series of the last 20 polling ticks.",
    "**Cell Type Distribution** — breakdown by Macro/Micro/Pico/Femto with health bars.",
]:
    bullet(doc, item)
figure_placeholder(doc, 18,
    "Network Topology with hover tooltip — cell type, location zone in cyan, health status, and KPI table.\n"
    "Screenshot: http://localhost:4000 (hover a critical/warning cell)")

h2(doc, "10.2  Intent AI (/intent)")
body(doc,
    "The primary interface for submitting natural language intents and reviewing AI pipeline results. "
    "The left panel contains the submission form with sample intents and history. "
    "The right panel renders `IntentResultCard` in six sections:")
for i, item in enumerate([
    "Intent Banner — intent type, slice, zone, confidence, health score, action.",
    "Agent Pipeline Timeline — four steps with checkmarks and summaries.",
    "KPI Before/After Comparison — four metric cards with % change and Improved/Degraded labels.",
    "Network Configuration Table — all 3GPP parameters grouped by category.",
    "Network Monitor Details — raw cell KPIs and threshold violations.",
    "Affected Topology — zone-highlighted SVG with monitored cell in cyan.",
], 1):
    numbered(doc, f"**{item.split('—')[0].strip()}** —{item.split('—')[1]}")
figure_placeholder(doc, 19,
    "Intent AI Page — full two-panel layout after a successful intent submission.\n"
    "Screenshot: http://localhost:4000/intent")

h2(doc, "10.3  Dataset Manager (/dataset)")
body(doc,
    "Allows engineers to inspect and replace the dataset. Displays filename, row count, cell count, "
    "source, and column list. Custom CSV uploads are validated client-side against all 11 required "
    "columns before acceptance. Engineers can revert to the default dataset at any time.")
figure_placeholder(doc, 20,
    "Dataset validation error — upload a CSV missing Location_Tag to demonstrate the rejection message.\n"
    "Screenshot: http://localhost:4000/dataset")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — Deployment
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "11. Deployment")

h2(doc, "11.1  Backend — Railway")
for step in [
    "Push `ran-dashboard-backend/` to a GitHub repository.",
    "In Railway: **New Project** → **Deploy from GitHub repo**.",
    "Set Root Directory to `ran-dashboard-backend/` if inside a monorepo.",
    "Add environment variable: `GROQ_API_KEY = gsk_your_key_here`.",
    "Railway auto-detects the Procfile and deploys automatically.",
]:
    numbered(doc, step)
code_block(doc, "web: gunicorn server:app --bind 0.0.0.0:$PORT --workers 1 --timeout 120")
note(doc, "The --timeout 120 flag is essential. LLM calls take 20–40 s and exceed gunicorn's default 30 s worker timeout.")

h2(doc, "11.2  Frontend — Vercel")
for step in [
    "Push `frontend/` to a GitHub repository.",
    "In Vercel: **New Project** → connect your GitHub repository.",
    "Set Root Directory to `frontend/` if deploying from a monorepo.",
    "Add environment variable: `NEXT_PUBLIC_BACKEND_URL=https://your-railway-app.up.railway.app`.",
    "Deploy — Vercel auto-detects Next.js. Every Git push triggers automatic redeployment.",
]:
    numbered(doc, step)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — Environment Variables
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "12. Environment Variables")

add_table(doc,
    ["Variable", "Where", "Required", "Description"],
    [
        ["GROQ_API_KEY",              "Backend",  "Yes",      "Groq API key. Without it all requests fall back to rule-based engine."],
        ["NEXT_PUBLIC_BACKEND_URL",   "Frontend", "Optional", "Deployed backend URL. Must be set on Vercel for production."],
        ["PORT",                      "Backend",  "Auto",     "Set automatically by Railway. Do not set manually in production."],
        ["FLASK_ENV",                 "Backend",  "Optional", "Set to 'development' for debug mode. Never use in production."],
    ],
    col_widths=[4.5, 2.5, 2.5, 6.0],
)

# ═══════════════════════════════════════════════════════════════════════════════
# Figure Index
# ═══════════════════════════════════════════════════════════════════════════════
h1(doc, "Figure Index")

add_table(doc,
    ["Fig.", "Description", "Source"],
    [
        ["1",  "Dashboard overview — Network Monitor page",               "Screenshot: localhost:4000"],
        ["2",  "5G three-slice architecture (eMBB / URLLC / mMTC)",       "Diagram"],
        ["3",  "5G RAN architecture — UE → gNB → 5G Core",                "Diagram"],
        ["4",  "HetNet topology — multi-tier cells with hover tooltip",    "Screenshot: localhost:4000"],
        ["5",  "KPI Summary Cards — live network averages",                "Screenshot: localhost:4000"],
        ["6",  "Live KPI Time-Series Chart",                               "Screenshot: localhost:4000"],
        ["7",  "IBN workflow diagram",                                      "Diagram"],
        ["8",  "Intent submission form",                                    "Screenshot: localhost:4000/intent"],
        ["9",  "Agent Pipeline Timeline",                                   "Screenshot: localhost:4000/intent"],
        ["10", "End-to-end data flow diagram",                              "Diagram"],
        ["11", "Location zone map — 5-zone spatial partitioning",           "Diagram"],
        ["12", "Dataset Manager page",                                      "Screenshot: localhost:4000/dataset"],
        ["13", "Full Intent Result Card",                                   "Screenshot: localhost:4000/intent"],
        ["14", "Intent Summary Banner",                                     "Screenshot: close-up"],
        ["15", "Generated 3GPP Network Configuration Table",               "Screenshot: close-up"],
        ["16", "KPI Before/After Comparison cards",                         "Screenshot: close-up"],
        ["17", "Affected Topology — zone-highlighted cells",                "Screenshot: close-up"],
        ["18", "Network Topology with hover tooltip",                       "Screenshot: localhost:4000"],
        ["19", "Intent AI Page — full two-panel layout",                    "Screenshot: localhost:4000/intent"],
        ["20", "Dataset validation error message",                          "Screenshot: localhost:4000/dataset"],
    ],
    col_widths=[1.5, 9.0, 5.0],
)

# ═══════════════════════════════════════════════════════════════════════════════
# References
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, "References")

refs = [
    '[1] 3rd Generation Partnership Project (3GPP), "System Architecture for the 5G System (5GS)," Technical Specification TS 23.501, Release 18, Dec. 2023.',
    '[2] 3rd Generation Partnership Project (3GPP), "Intent driven management services for mobile networks; Requirements," Technical Specification TS 28.312, Release 18, 2023.',
    '[3] A. Clemm, L. Ciavaglia, L. Z. Granville, and J. Tantsura, "Intent-Based Networking — Concepts and Definitions," IETF, Request for Comments RFC 9315, Oct. 2022.',
    '[4] J. G. Andrews, S. Buzzi, W. Choi, S. V. Hanly, A. Lozano, A. C. K. Soong, and J. C. Zhang, "What Will 5G Be?" IEEE Journal on Selected Areas in Communications, vol. 32, no. 6, pp. 1065–1082, Jun. 2014. DOI: 10.1109/JSAC.2014.2328098.',
    '[5] A. Damnjanovic, J. Montojo, Y. Wei, T. Ji, T. Luo, M. Vajapeyam, T. Yoo, O. Song, and D. Malladi, "A Survey on 3GPP Heterogeneous Networks," IEEE Wireless Communications, vol. 18, no. 3, pp. 10–21, Jun. 2011. DOI: 10.1109/MWC.2011.5876496.',
    '[6] AI @ Meta, "The Llama 3 Herd of Models," arXiv preprint arXiv:2407.21783, Jul. 2024.',
    '[7] M. Polese, L. Bonati, S. D\'Oro, S. Basagni, and T. Melodia, "Understanding O-RAN: Architecture, Interfaces, Algorithms, Security, and Research Challenges," IEEE Communications Surveys & Tutorials, vol. 25, no. 2, pp. 1376–1411, 2023. DOI: 10.1109/COMST.2023.3239220.',
    '[8] X. Foukas, G. Patounas, A. Elmokashfi, and M. K. Marina, "Network Slicing in 5G: Survey and Challenges," IEEE Communications Magazine, vol. 55, no. 5, pp. 94–100, May 2017. DOI: 10.1109/MCOM.2017.1600951.',
    '[9] R. Li, Z. Zhao, X. Zhou, G. Ding, Y. Chen, Z. Wang, and H. Zhang, "Intelligent 5G: When Cellular Networks Meet Artificial Intelligence," IEEE Wireless Communications, vol. 24, no. 5, pp. 175–183, Oct. 2017. DOI: 10.1109/MWC.2017.1600304WC.',
    '[10] J. Wei, X. Wang, D. Schuurmans, M. Bosma, F. Xia, E. Chi, Q. V. Le, and D. Zhou, "Chain-of-Thought Prompting Elicits Reasoning in Large Language Models," in Advances in Neural Information Processing Systems (NeurIPS), vol. 35, pp. 24824–24837, 2022.',
    '[11] 3rd Generation Partnership Project (3GPP), "5G; NR; Physical Layer Procedures for Data," Technical Specification TS 38.214, Release 18, 2023.',
    '[12] TMForum, "Autonomous Networks: Empowering Digital Transformation," Technical Report TR290, Version 2.0, 2021.',
    '[13] A. Kaloxylos, "A Survey and an Analysis of Network Slicing in 5G Networks," IEEE Communications Standards Magazine, vol. 2, no. 1, pp. 60–65, Mar. 2018. DOI: 10.1109/MCOMSTD.2018.1700072.',
    '[14] S. Andreev et al., "Understanding the IoT Connectivity Landscape: A Contemporary M2M Radio Technology Roadmap," IEEE Communications Magazine, vol. 53, no. 9, pp. 32–40, Sep. 2015. DOI: 10.1109/MCOM.2015.7263370.',
    '[15] CrewAI, Inc., "CrewAI: Framework for Orchestrating Role-Playing, Autonomous AI Agents," GitHub Repository, 2024. [Online]. Available: https://github.com/crewAIInc/crewAI',
    '[16] Groq, Inc., "Groq LPU Inference Engine," Technical Overview, 2024. [Online]. Available: https://groq.com',
    '[17] T. Wolf et al., "Transformers: State-of-the-Art Natural Language Processing," in Proc. EMNLP System Demonstrations, pp. 38–45, 2020. DOI: 10.18653/v1/2020.emnlp-demos.6.',
    '[18] Vercel, Inc., "Next.js 14 Documentation," 2024. [Online]. Available: https://nextjs.org/docs',
    '[19] W. McKinney, "Data Structures for Statistical Computing in Python," in Proc. 9th Python in Science Conference (SciPy), pp. 56–61, 2010.',
    '[20] ETSI, "Experiential Networked Intelligence (ENI); ENI Architecture," ETSI GS ENI 005, V2.1.1, Jun. 2020.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.space_after       = Pt(4)
    r = p.add_run(ref)
    r.font.name = "Calibri"
    r.font.size = Pt(10)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\danyh\ran-dashboard-backend\RAN_Optimization_Dashboard_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
