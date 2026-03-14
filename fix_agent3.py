"""
Fix the Word document:
1. Remove the wrongly placed Safety Validator content (paragraphs 58-60 + table)
2. Re-insert it correctly before "8.3 Agent 3 — Network Monitor"
3. Renumber 8.3→8.4 (Monitor) and 8.4→8.5 (Optimizer)
4. Update "Four-Agent" → "Five-Agent" in section heading and other references
Run: python fix_agent3.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_PATH = "RAN_Optimization_Dashboard_Documentation.docx"
doc = Document(DOC_PATH)
body_el = doc.element.body

# ── Step 1: Remove the wrong insertion (paras 58, 59, 60 + table between them) ─

paras = doc.paragraphs
wrong_texts = [
    "A rejected verdict means the configuration",
    "Role: RAN Configuration Safety Validator. Validates",
    "Agent 3: Safety Validator",
]

for wrong in wrong_texts:
    for p in doc.paragraphs:
        if p.text.strip().startswith(wrong):
            p._element.getparent().remove(p._element)
            print(f"Removed paragraph: {wrong[:50]}")
            break

# Remove wrongly placed table (find table whose first cell contains safety text)
for tbl in doc.tables:
    try:
        first_cell = tbl.rows[0].cells[0].text.strip()
        if first_cell in ("Check", "Bandwidth limit"):
            tbl._element.getparent().remove(tbl._element)
            print("Removed wrongly placed table")
            break
    except Exception:
        pass

# ── Step 2: Update text references ─────────────────────────────────────────────

# Two separate passes to avoid double-rename (8.3->8.4->8.5 on same run)
replacements_pass1 = {
    "The Four-Agent Crew":    "The Five-Agent Crew",
    "four-agent pipeline":    "five-agent pipeline",
    "full four-agent":        "full five-agent",
    "four agent outputs":     "five agent outputs",
    "Agents 3 and 4 call":    "Agents 4 and 5 call",
    "8.4  Agent 4":           "8.5  Agent 5",       # rename old 4 first
    "Agent 4 \u2013 RAN Opt": "Agent 5 \u2013 RAN Opt",
}
replacements_pass2 = {
    "8.3  Agent 3":           "8.4  Agent 4",       # then rename old 3
    "Agent 3 \u2013 Network": "Agent 4 \u2013 Network",
}

def apply_replacements(replacements):
    for p in doc.paragraphs:
        for run in p.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    print(f"Updated: {old[:40]}")
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    for run in cp.runs:
                        for old, new in replacements.items():
                            if old in run.text:
                                run.text = run.text.replace(old, new)

apply_replacements(replacements_pass1)
apply_replacements(replacements_pass2)

# ── Style helpers ─────────────────────────────────────────────────────────────

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, color="BDD7EE"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def make_cell_text(cell, text, bold=False, color_rgb=None, font_size=10):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

# ── Step 3: Build new content in a temp doc ───────────────────────────────────

tmp = Document()

# Heading
h3_p = tmp.add_paragraph()
h3_p.paragraph_format.space_before = Pt(8)
h3_p.paragraph_format.space_after  = Pt(3)
r = h3_p.add_run("8.3  Agent 3 \u2013 Safety Validator")
r.font.name = "Calibri"; r.font.size = Pt(11.5); r.font.bold = True
r.font.color.rgb = RGBColor(0x40, 0x60, 0x8A)

# Body intro
body_p = tmp.add_paragraph()
body_p.paragraph_format.space_before = Pt(0)
body_p.paragraph_format.space_after  = Pt(6)
r = body_p.add_run(
    "Role: RAN Configuration Safety Validator. Validates the RAN Planner\u2019s output "
    "against four hard limits before any configuration can proceed to monitoring or execution."
)
r.font.name = "Calibri"; r.font.size = Pt(11)

# Table
tbl = tmp.add_table(rows=5, cols=3)
tbl.style = "Table Grid"

headers = ["Check", "Limit", "Verdict on Failure"]
hdr = tbl.rows[0]
for i, h in enumerate(headers):
    shade_cell(hdr.cells[i], "2B4C7E")
    set_cell_border(hdr.cells[i], color="2B4C7E")
    make_cell_text(hdr.cells[i], h, bold=True, color_rgb=(0xFF,0xFF,0xFF))

data = [
    ("Bandwidth limit",     "\u2264 1000 Mbps per slice",                         "Rejected"),
    ("Latency feasibility", "URLLC \u2264 10 ms \u00b7 eMBB \u2264 100 ms \u00b7 mMTC \u2264 300 ms", "Rejected"),
    ("Capacity check",      "expected_users \u2264 active_cells \u00d7 1000",      "Approved with warnings"),
    ("QoS consistency",     "5QI value matches slice type",                       "Approved with warnings"),
]
for ri, (c1, c2, c3) in enumerate(data):
    row = tbl.rows[ri + 1]
    fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate((c1, c2, c3)):
        shade_cell(row.cells[ci], fill)
        set_cell_border(row.cells[ci])
        make_cell_text(row.cells[ci], val)

# Closing note
note_p = tmp.add_paragraph()
note_p.paragraph_format.space_before = Pt(4)
note_p.paragraph_format.space_after  = Pt(6)
r = note_p.add_run(
    "A rejected verdict means the configuration violates a hard physical or standards "
    "constraint and cannot be deployed. An approved with warnings verdict means the "
    "configuration is deployable but the operator is alerted to a potential issue."
)
r.font.name = "Calibri"; r.font.size = Pt(11)

# ── Step 4: Find correct insertion point and insert ───────────────────────────

# Find "8.4  Agent 4 — Network Monitor" (already renamed above)
target_p_el = None
for p in doc.paragraphs:
    if "8.4  Agent 4" in p.text and "Network" in p.text:
        target_p_el = p._p
        print(f"Insertion point found: {p.text[:60]}")
        break

if target_p_el is None:
    # fallback: try original name if rename didn't fire
    for p in doc.paragraphs:
        if "8.3  Agent 3" in p.text and "Network" in p.text:
            target_p_el = p._p
            print(f"Insertion point (fallback): {p.text[:60]}")
            break

if target_p_el is None:
    print("ERROR: Could not find insertion point!")
    exit(1)

# Insert in correct order (forward, each addprevious before target)
new_elements = [h3_p._p, body_p._p, tbl._tbl, note_p._p]
for el in new_elements:
    el_copy = copy.deepcopy(el)
    target_p_el.addprevious(el_copy)

print("Inserted Agent 3: Safety Validator section in correct location.")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")
